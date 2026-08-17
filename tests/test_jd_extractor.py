"""Unit tests for JobDescription model and RuleBasedJobDescriptionExtractor."""

from pathlib import Path
import docx
import pytest

from app.extractors import RuleBasedJobDescriptionExtractor, extract_job_description
from app.models import JobDescription


@pytest.fixture
def sample_jd_text() -> str:
    """Realistic standard job description text."""
    return """
    Job Title: Junior AI Research Associate

    RESPONSIBILITIES
    - Develop, train, and evaluate deep learning and NLP models.
    - Assist in designing benchmarking pipelines and evaluation metrics.
    - Collaborate with the research engineering team to deploy inference APIs.

    REQUIREMENTS
    - Bachelor's Degree or Master's Degree in Computer Science, Data Science, or related field.
    - Minimum 1+ years of experience with machine learning and software development.
    - Strong proficiency in Python, PyTorch, and Scikit-Learn.
    - Solid understanding of Git and Linux.

    PREFERRED QUALIFICATIONS
    - Experience with Docker, Kubernetes, and FastAPI.
    - Familiarity with Large Language Models, Transformers, and BERT.
    """


@pytest.fixture
def jd_extractor() -> RuleBasedJobDescriptionExtractor:
    """Provide RuleBasedJobDescriptionExtractor instance."""
    return RuleBasedJobDescriptionExtractor()


class TestJobDescriptionModel:
    """Test suite for JobDescription Pydantic model validation."""

    def test_jd_defaults(self):
        """Test default values for JobDescription model."""
        jd = JobDescription()
        assert jd.title == "Unknown Position"
        assert jd.required_skills == []
        assert jd.preferred_skills == []
        assert jd.minimum_experience_years is None
        assert jd.education_requirements == []
        assert jd.responsibilities == []
        assert jd.raw_text == ""

    def test_skill_deduplication(self):
        """Test that skills are deduplicated case-insensitively."""
        jd = JobDescription(
            title="AI Engineer",
            required_skills=["Python", "python", "PyTorch", "pytorch"],
            preferred_skills=["Docker", "docker", "AWS"],
        )
        assert jd.required_skills == ["Python", "PyTorch"]
        assert jd.preferred_skills == ["Docker", "AWS"]


class TestJobDescriptionExtractor:
    """Test suite for RuleBasedJobDescriptionExtractor."""

    def test_extract_normal_jd(
        self, jd_extractor: RuleBasedJobDescriptionExtractor, sample_jd_text: str
    ):
        """Test extracting all fields from a complete, structured job description."""
        jd = jd_extractor.extract_from_text(sample_jd_text)

        assert jd.title == "Junior AI Research Associate"
        assert "Python" in jd.required_skills
        assert "PyTorch" in jd.required_skills
        assert "Scikit-Learn" in jd.required_skills
        assert "Git" in jd.required_skills
        assert "Linux" in jd.required_skills

        assert "Docker" in jd.preferred_skills
        assert "FastAPI" in jd.preferred_skills
        assert "Large Language Models" in jd.preferred_skills

        assert jd.minimum_experience_years == 1.0
        assert "Bachelor's Degree" in jd.education_requirements
        assert "Master's Degree" in jd.education_requirements
        assert "Degree in Computer Science" in jd.education_requirements

        assert len(jd.responsibilities) == 3
        assert any("deep learning" in r.lower() for r in jd.responsibilities)

    def test_required_vs_preferred_segregation(
        self, jd_extractor: RuleBasedJobDescriptionExtractor
    ):
        """Test that required and preferred skills are not conflated."""
        text = """
        Position: Backend Engineer

        REQUIREMENTS
        - Python, PostgreSQL, SQL

        NICE TO HAVE
        - Docker, AWS, Kubernetes
        """
        jd = jd_extractor.extract_from_text(text)
        assert "Python" in jd.required_skills
        assert "PostgreSQL" in jd.required_skills
        assert "SQL" in jd.required_skills

        assert "Docker" in jd.preferred_skills
        assert "AWS" in jd.preferred_skills
        assert "Kubernetes" in jd.preferred_skills

        # Ensure no overlap
        for req in jd.required_skills:
            assert req not in jd.preferred_skills

    def test_experience_extraction_variations(
        self, jd_extractor: RuleBasedJobDescriptionExtractor
    ):
        """Test various experience requirement phrases."""
        text_1 = "Requirements: Minimum 3 years of software experience."
        jd_1 = jd_extractor.extract_from_text(text_1)
        assert jd_1.minimum_experience_years == 3.0

        text_2 = "Requirements: 2+ years of hands-on machine learning experience."
        jd_2 = jd_extractor.extract_from_text(text_2)
        assert jd_2.minimum_experience_years == 2.0

        text_3 = "Requirements: At least 5 years experience."
        jd_3 = jd_extractor.extract_from_text(text_3)
        assert jd_3.minimum_experience_years == 5.0

    def test_education_requirements(
        self, jd_extractor: RuleBasedJobDescriptionExtractor
    ):
        """Test degree and discipline extraction."""
        text = """
        Role: Senior Research Scientist
        QUALIFICATIONS
        - Ph.D. in Artificial Intelligence or Robotics required.
        """
        jd = jd_extractor.extract_from_text(text)
        assert "Ph.D." in jd.education_requirements
        assert "Degree in Artificial Intelligence" in jd.education_requirements
        assert "Degree in Robotics" in jd.education_requirements

    def test_responsibilities_extraction(
        self, jd_extractor: RuleBasedJobDescriptionExtractor
    ):
        """Test extracting role responsibilities."""
        text = """
        Role: ML Engineer
        RESPONSIBILITIES
        • Build end-to-end data processing pipelines
        • Optimize inference latency for real-time models
        • Maintain CI/CD pipelines
        """
        jd = jd_extractor.extract_from_text(text)
        assert len(jd.responsibilities) == 3
        assert "Build end-to-end data processing pipelines" in jd.responsibilities
        assert "Optimize inference latency for real-time models" in jd.responsibilities

    def test_missing_sections(
        self, jd_extractor: RuleBasedJobDescriptionExtractor
    ):
        """Test that missing sections cleanly default without crashing."""
        text = "We are hiring a Python developer with knowledge of FastAPI."
        jd = jd_extractor.extract_from_text(text)
        assert "Python" in jd.required_skills
        assert "FastAPI" in jd.preferred_skills or "FastAPI" in jd.required_skills
        assert jd.preferred_skills == [] or "FastAPI" in jd.preferred_skills
        assert jd.minimum_experience_years is None
        assert jd.education_requirements == []
        assert jd.responsibilities == []

    def test_skill_normalization(
        self, jd_extractor: RuleBasedJobDescriptionExtractor
    ):
        """Test that skill aliases are normalized using shared dictionary."""
        text = """
        REQUIREMENTS
        - python3, pytorch, tf, sklearn, ml
        """
        jd = jd_extractor.extract_from_text(text)
        assert "Python" in jd.required_skills
        assert "PyTorch" in jd.required_skills
        assert "TensorFlow" in jd.required_skills
        assert "Scikit-Learn" in jd.required_skills
        assert "Machine Learning" in jd.required_skills

    def test_extract_from_txt_file(
        self, tmp_path: Path, sample_jd_text: str
    ):
        """Test end-to-end extraction from a TXT file on disk."""
        jd_file = tmp_path / "job_description.txt"
        jd_file.write_text(sample_jd_text, encoding="utf-8")

        jd = extract_job_description(jd_file)
        assert jd.title == "Junior AI Research Associate"
        assert "Python" in jd.required_skills
        assert jd.minimum_experience_years == 1.0

    def test_extract_from_docx_file(self, tmp_path: Path):
        """Test end-to-end extraction from a DOCX job description file."""
        docx_file = tmp_path / "jd.docx"
        doc = docx.Document()
        doc.add_heading("Job Title: Data Scientist", level=1)
        doc.add_paragraph("REQUIREMENTS")
        doc.add_paragraph("- Python, Pandas, SQL")
        doc.add_paragraph("- 2+ years of experience")
        doc.save(str(docx_file))

        jd = extract_job_description(docx_file)
        assert jd.title == "Data Scientist"
        assert "Python" in jd.required_skills
        assert "Pandas" in jd.required_skills
        assert "SQL" in jd.required_skills
        assert jd.minimum_experience_years == 2.0

    def test_malformed_and_empty_text(
        self, jd_extractor: RuleBasedJobDescriptionExtractor
    ):
        """Test behavior with empty or whitespace-only JD text."""
        empty_jd = jd_extractor.extract_from_text("")
        assert empty_jd.title == "Unknown Position"
        assert empty_jd.required_skills == []
        assert empty_jd.preferred_skills == []
        assert empty_jd.minimum_experience_years is None

        ws_jd = jd_extractor.extract_from_text("   \n\n\t  ")
        assert ws_jd.title == "Unknown Position"
        assert ws_jd.required_skills == []
