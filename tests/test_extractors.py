"""Unit tests for Candidate model and Resume Extractor."""

from pathlib import Path
import docx
import pytest

from app.extractors import RuleBasedResumeExtractor, extract_resume
from app.models import Candidate, Education, Experience


@pytest.fixture
def sample_resume_text() -> str:
    """Realistic complete resume text."""
    return """
    Alex Johnson
    Email: alex.johnson@example.com
    Phone: +1 (555) 123-4567
    GitHub: github.com/alexjohnson

    SUMMARY
    Junior AI Researcher passionate about machine learning, NLP, and deep learning architectures.

    SKILLS
    Python, Python3, PyTorch, PyTorch, TensorFlow, Scikit-Learn, Pandas, Git, Docker, C++

    EDUCATION
    Master of Science in Computer Science
    Stanford University, 2023

    EXPERIENCE
    AI Research Assistant | Stanford AI Lab
    2021 - 2023
    - Conducted transformer interpretability research
    - Trained benchmark vision-language models
    """


@pytest.fixture
def extractor() -> RuleBasedResumeExtractor:
    """Provide RuleBasedResumeExtractor instance."""
    return RuleBasedResumeExtractor()


class TestCandidateModel:
    """Test suite for Candidate Pydantic model validation."""

    def test_candidate_defaults(self):
        """Test default values for Candidate model."""
        candidate = Candidate()
        assert candidate.name == "Unknown Candidate"
        assert candidate.email is None
        assert candidate.phone is None
        assert candidate.skills == []
        assert candidate.education == []
        assert candidate.experience == []

    def test_skill_deduplication(self):
        """Test that skill list is deduplicated case-insensitively while preserving formatting."""
        candidate = Candidate(
            name="Jane Doe",
            skills=["Python", "python", "PyTorch", "pytorch", "Docker"],
        )
        assert len(candidate.skills) == 3
        assert "Python" in candidate.skills
        assert "PyTorch" in candidate.skills
        assert "Docker" in candidate.skills


class TestResumeExtractor:
    """Test suite for RuleBasedResumeExtractor."""

    def test_extract_normal_resume(self, extractor: RuleBasedResumeExtractor, sample_resume_text: str):
        """Test extraction on a standard resume with all fields present."""
        candidate = extractor.extract_from_text(sample_resume_text)

        assert candidate.name == "Alex Johnson"
        assert candidate.email == "alex.johnson@example.com"
        assert candidate.phone == "+1 (555) 123-4567"
        assert "Python" in candidate.skills
        assert "PyTorch" in candidate.skills
        assert "TensorFlow" in candidate.skills
        assert "Docker" in candidate.skills

        # Education
        assert len(candidate.education) >= 1
        edu = candidate.education[0]
        assert edu.degree == "Master's Degree"
        assert edu.field_of_study == "Computer Science"
        assert edu.end_year == 2023
        assert "Stanford University" in edu.institution

        # Experience
        assert len(candidate.experience) >= 1
        exp = candidate.experience[0]
        assert exp.start_date == "2021"
        assert exp.end_date == "2023"
        assert "Research" in exp.title
        assert exp.company == "Stanford AI Lab"

    def test_missing_email(self, extractor: RuleBasedResumeExtractor):
        """Test resume without email address."""
        text = """
        John Smith
        Phone: 555-987-6543
        SKILLS: Python, SQL
        """
        candidate = extractor.extract_from_text(text)
        assert candidate.name == "John Smith"
        assert candidate.email is None
        assert candidate.phone == "555-987-6543"
        assert "Python" in candidate.skills

    def test_missing_phone(self, extractor: RuleBasedResumeExtractor):
        """Test resume without phone number."""
        text = """
        John Smith
        Email: jsmith@example.org
        SKILLS: Machine Learning, Deep Learning
        """
        candidate = extractor.extract_from_text(text)
        assert candidate.name == "John Smith"
        assert candidate.email == "jsmith@example.org"
        assert candidate.phone is None
        assert "Machine Learning" in candidate.skills
        assert "Deep Learning" in candidate.skills

    def test_skill_normalization_and_deduplication(self, extractor: RuleBasedResumeExtractor):
        """Test that aliases (python3, py, torch, ml) normalize to canonical names without duplicates."""
        text = """
        Candidate: Priya Sharma
        Email: priya@example.com
        SKILLS: python3, python, pytorch, torch, ml, machine learning, natural language processing, nlp
        """
        candidate = extractor.extract_from_text(text)
        assert "Python" in candidate.skills
        assert "PyTorch" in candidate.skills
        assert "Machine Learning" in candidate.skills
        assert "Natural Language Processing" in candidate.skills
        # Ensure no duplicates
        assert len([s for s in candidate.skills if s == "Python"]) == 1
        assert len([s for s in candidate.skills if s == "PyTorch"]) == 1

    def test_education_variations(self, extractor: RuleBasedResumeExtractor):
        """Test extracting Ph.D., B.S., B.Tech variations."""
        text = """
        Dr. Robert King
        Email: rking@lab.edu
        EDUCATION
        Ph.D. in Artificial Intelligence
        MIT, 2020
        B.Tech in Electrical Engineering, 2015
        """
        candidate = extractor.extract_from_text(text)
        assert len(candidate.education) == 2
        degrees = [e.degree for e in candidate.education]
        assert "Ph.D." in degrees
        assert "Bachelor's Degree" in degrees

    def test_education_year_range_regression(self, extractor: RuleBasedResumeExtractor):
        """Regression test for education date ranges (start_year and end_year)."""
        text = """
        ADITYA SHARMA
        EDUCATION
        B.Tech in Computer Science Engineering
        ABC Institute of Technology
        2022 - 2026
        """
        candidate = extractor.extract_from_text(text)
        assert len(candidate.education) == 1
        edu = candidate.education[0]
        assert edu.degree == "Bachelor's Degree"
        assert edu.field_of_study == "Computer Science"
        assert edu.institution == "ABC Institute of Technology"
        assert edu.start_year == 2022
        assert edu.end_year == 2026

    def test_experience_company_structural_heuristic(self, extractor: RuleBasedResumeExtractor):
        """Regression test for multi-line Title -> Company -> Dates -> Description structure."""
        text = """
        EXPERIENCE
        Software Engineering Intern
        Tech Solutions Pvt Ltd
        Jan 2025 - Jun 2025
        Built REST APIs using Python and FastAPI.
        Worked with PostgreSQL and Docker.
        """
        candidate = extractor.extract_from_text(text)
        assert len(candidate.experience) == 1
        exp = candidate.experience[0]
        assert exp.title == "Software Engineering Intern"
        assert exp.company == "Tech Solutions Pvt Ltd"
        assert exp.start_date == "Jan 2025"
        assert exp.end_date == "Jun 2025"
        assert "Built REST APIs" in exp.description
        assert "Tech Solutions Pvt Ltd" not in exp.description

    def test_experience_extraction(self, extractor: RuleBasedResumeExtractor):
        """Test extracting structured experience with date ranges and descriptions."""
        text = """
        Sarah Connor
        Email: sconnor@tech.co
        EXPERIENCE
        Software Engineer | Cyberdyne Systems
        Jan 2020 - Present
        - Developed autonomous agents
        - Managed Kubernetes clusters
        """
        candidate = extractor.extract_from_text(text)
        assert len(candidate.experience) >= 1
        exp = candidate.experience[0]
        assert "Engineer" in exp.title
        assert exp.company == "Cyberdyne Systems"
        assert "Jan 2020" in exp.start_date
        assert "Present" in exp.end_date

    def test_empty_and_malformed_text(self, extractor: RuleBasedResumeExtractor):
        """Test extractor behavior with empty or minimal text."""
        candidate_empty = extractor.extract_from_text("")
        assert candidate_empty.name == "Unknown Candidate"
        assert candidate_empty.email is None
        assert candidate_empty.skills == []

        candidate_whitespace = extractor.extract_from_text("   \n\n   ")
        assert candidate_whitespace.name == "Unknown Candidate"
        assert candidate_whitespace.skills == []

    def test_extract_from_txt_file(self, tmp_path: Path, sample_resume_text: str):
        """Test end-to-end extraction from a TXT file on disk using extract_resume()."""
        resume_file = tmp_path / "alex_resume.txt"
        resume_file.write_text(sample_resume_text, encoding="utf-8")

        candidate = extract_resume(resume_file)
        assert candidate.name == "Alex Johnson"
        assert candidate.email == "alex.johnson@example.com"
        assert "Python" in candidate.skills
        assert len(candidate.education) >= 1

    def test_extract_from_docx_file(self, tmp_path: Path):
        """Test end-to-end extraction from a DOCX resume file."""
        docx_file = tmp_path / "candidate.docx"
        doc = docx.Document()
        doc.add_paragraph("Elena Rostova")
        doc.add_paragraph("Email: elena@ai-research.org")
        doc.add_paragraph("Phone: 555-321-9876")
        doc.add_paragraph("SKILLS")
        doc.add_paragraph("PyTorch, Transformers, Deep Learning, Linux")
        doc.add_paragraph("EDUCATION")
        doc.add_paragraph("Ph.D. in Computer Science")
        doc.add_paragraph("Oxford University, 2022")
        doc.save(str(docx_file))

        candidate = extract_resume(docx_file)
        assert candidate.name == "Elena Rostova"
        assert candidate.email == "elena@ai-research.org"
        assert candidate.phone == "555-321-9876"
        assert "PyTorch" in candidate.skills
        assert "Transformers" in candidate.skills
        assert candidate.education[0].degree == "Ph.D."

    def test_sample_resume_file_integration(self):
        """End-to-end integration test on data/resumes/sample_resume.txt."""
        sample_path = Path("data/resumes/sample_resume.txt")
        if sample_path.exists():
            candidate = extract_resume(sample_path)
            assert candidate.name == "ADITYA SHARMA"
            assert candidate.email == "aditya.sharma@email.com"
            assert candidate.phone == "+91 9876543210"
            assert "Python" in candidate.skills
            assert "FastAPI" in candidate.skills
            assert "PostgreSQL" in candidate.skills
            assert len(candidate.education) == 1
            assert candidate.education[0].start_year == 2022
            assert candidate.education[0].end_year == 2026
            assert len(candidate.experience) == 1
            assert candidate.experience[0].company == "Tech Solutions Pvt Ltd"
            assert candidate.experience[0].start_date == "Jan 2025"
            assert candidate.experience[0].end_date == "Jun 2025"
