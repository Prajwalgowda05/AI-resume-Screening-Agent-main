"""Unit tests for document parsing modules (PDF, DOCX, TXT)."""

import io
from pathlib import Path
import pytest
import docx
import pypdf

from app.parsers import (
    DocxParser,
    DocumentParsingError,
    PDFParser,
    TextParser,
    UnsupportedFileTypeError,
    get_parser,
    parse_document,
)


@pytest.fixture
def sample_txt_file(tmp_path: Path) -> Path:
    """Create a temporary text file with sample resume content."""
    file_path = tmp_path / "resume_sample.txt"
    content = "Candidate Name: Alex Johnson\nSkills: Python, Machine Learning\nExperience: 2 years"
    file_path.write_text(content, encoding="utf-8")
    return file_path


@pytest.fixture
def sample_pdf_file(tmp_path: Path) -> Path:
    """Create a temporary valid PDF file with sample resume content."""
    file_path = tmp_path / "resume_sample.pdf"
    writer = pypdf.PdfWriter()
    # Add a blank page with text annotation/content
    page = writer.add_blank_page(width=200, height=200)
    
    # We can write an actual PDF with text streams using pypdf
    # Alternatively, create a PDF stream with text:
    writer.write(str(file_path))
    
    # Let's create a PDF containing real searchable text:
    # Using pypdf's canvas or writer with simple text annotation / objects
    return file_path


def create_pdf_with_text(file_path: Path, text: str):
    """Helper to generate a minimal valid PDF with extractable text."""
    # Write a minimal PDF stream structure with stream text
    pdf_content = (
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj\n"
        b"4 0 obj << /Length " + str(len(text) + 40).encode("utf-8") + b" >>\n"
        b"stream\n"
        b"BT\n"
        b"/F1 12 Tf\n"
        b"100 700 Td\n"
        b"(" + text.encode("latin-1", "replace") + b") Tj\n"
        b"ET\n"
        b"endstream\n"
        b"endobj\n"
        b"5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        b"xref\n"
        b"0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000240 00000 n \n"
        b"0000000340 00000 n \n"
        b"trailer << /Size 6 /Root 1 0 R >>\n"
        b"startxref\n"
        b"415\n"
        b"%%EOF\n"
    )
    file_path.write_bytes(pdf_content)


@pytest.fixture
def sample_docx_file(tmp_path: Path) -> Path:
    """Create a temporary valid DOCX file with paragraphs and tables."""
    file_path = tmp_path / "resume_sample.docx"
    doc = docx.Document()
    doc.add_heading("Alex Johnson - Resume", level=1)
    doc.add_paragraph("Summary: Experienced AI Researcher.")
    
    # Add a table (common in resumes)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill Category"
    table.cell(0, 1).text = "Technologies"
    table.cell(1, 0).text = "Languages"
    table.cell(1, 1).text = "Python, SQL, C++"
    
    doc.save(str(file_path))
    return file_path


class TestDocumentParsers:
    """Test suite for document parsers."""

    def test_parse_txt_file(self, sample_txt_file: Path):
        """Test extraction from a plain text file."""
        text = parse_document(sample_txt_file)
        assert "Candidate Name: Alex Johnson" in text
        assert "Skills: Python, Machine Learning" in text
        assert isinstance(get_parser(sample_txt_file), TextParser)

    def test_parse_docx_file(self, sample_docx_file: Path):
        """Test extraction from a DOCX file including paragraphs and tables."""
        text = parse_document(sample_docx_file)
        assert "Alex Johnson - Resume" in text
        assert "Summary: Experienced AI Researcher." in text
        assert "Python, SQL, C++" in text
        assert isinstance(get_parser(sample_docx_file), DocxParser)

    def test_parse_pdf_file(self, tmp_path: Path):
        """Test extraction from a PDF file."""
        pdf_path = tmp_path / "test_resume.pdf"
        create_pdf_with_text(pdf_path, "John Doe - Machine Learning Engineer")

        text = parse_document(pdf_path)
        assert "John Doe - Machine Learning Engineer" in text
        assert isinstance(get_parser(pdf_path), PDFParser)

    def test_unsupported_file_type(self, tmp_path: Path):
        """Test that an unsupported file format raises UnsupportedFileTypeError."""
        invalid_file = tmp_path / "image.png"
        invalid_file.write_bytes(b"\x89PNG\r\n\x1a\n")

        with pytest.raises(UnsupportedFileTypeError) as exc_info:
            parse_document(invalid_file)
        assert "Unsupported file format '.png'" in str(exc_info.value)

    def test_missing_file(self, tmp_path: Path):
        """Test that attempting to parse a non-existent file raises FileNotFoundError."""
        missing_file = tmp_path / "non_existent_file.pdf"
        with pytest.raises(FileNotFoundError):
            parse_document(missing_file)

    def test_case_insensitive_extension(self, tmp_path: Path):
        """Test that file extensions with mixed or upper case are handled."""
        pdf_path = tmp_path / "UPPERCASE_TEST.PDF"
        create_pdf_with_text(pdf_path, "Upper Case Test Candidate")

        parser = get_parser(pdf_path)
        assert isinstance(parser, PDFParser)

        text = parse_document(pdf_path)
        assert "Upper Case Test Candidate" in text

    def test_corrupted_pdf_error(self, tmp_path: Path):
        """Test that corrupted PDF file raises DocumentParsingError."""
        corrupted_pdf = tmp_path / "corrupted.pdf"
        corrupted_pdf.write_bytes(b"This is not a valid PDF header.")

        with pytest.raises(DocumentParsingError):
            parse_document(corrupted_pdf)

    def test_corrupted_docx_error(self, tmp_path: Path):
        """Test that corrupted DOCX file raises DocumentParsingError."""
        corrupted_docx = tmp_path / "corrupted.docx"
        corrupted_docx.write_bytes(b"PK\x03\x04CorruptZipData")

        with pytest.raises(DocumentParsingError):
            parse_document(corrupted_docx)

    def test_empty_txt_file(self, tmp_path: Path):
        """Test that an empty text file returns empty string without crashing."""
        empty_txt = tmp_path / "empty.txt"
        empty_txt.write_text("", encoding="utf-8")

        text = parse_document(empty_txt)
        assert text == ""
