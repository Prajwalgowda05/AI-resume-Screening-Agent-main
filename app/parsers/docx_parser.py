"""DOCX document parser implementation using python-docx."""

from pathlib import Path
from typing import Union
import docx

from app.parsers.base import BaseParser, DocumentParsingError


class DocxParser(BaseParser):
    """Parser for extracting text from DOCX files using python-docx."""

    def parse(self, file_path: Union[str, Path]) -> str:
        """Extract text from a DOCX file including paragraphs and table contents.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Extracted text content.

        Raises:
            FileNotFoundError: If the file does not exist.
            DocumentParsingError: If DOCX reading or extraction fails.
        """
        path = self.validate_file(file_path)

        try:
            doc = docx.Document(str(path))
            content_chunks = []

            # Extract body paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_chunks.append(text)

            # Extract text from tables (common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    ]
                    if row_cells:
                        # Deduplicate adjacent identical cells resulting from merged cells
                        unique_cells = []
                        for cell_text in row_cells:
                            if not unique_cells or unique_cells[-1] != cell_text:
                                unique_cells.append(cell_text)
                        content_chunks.append(" | ".join(unique_cells))

            return "\n".join(content_chunks).strip()
        except FileNotFoundError:
            raise
        except Exception as e:
            raise DocumentParsingError(
                f"Failed to extract text from DOCX '{path.name}': {str(e)}"
            ) from e
